import os
import io
from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Header, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
from supabase import create_client, Client

# Load environment variables from backend/.env relative to this file
env_path = Path(__file__).resolve().parent / '.env'
load_dotenv(dotenv_path=env_path)

from parser import extract_content, ExtractionError
from gemini import GeminiClient

app = FastAPI(title="TenderIQ API", version="1.0.0")

# ─── CORS ────────────────────────────────────────────────────────────────────
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Gemini Client ────────────────────────────────────────────────────────────
try:
    gemini_client = GeminiClient()
except Exception as e:
    print(f"WARNING: Gemini client could not be initialized: {e}")
    gemini_client = None


# ─── Supabase helper ──────────────────────────────────────────────────────────
def get_supabase(auth_header: Optional[str] = None) -> Client:
    url = os.environ.get("SUPABASE_URL") or os.environ.get("VITE_SUPABASE_URL")
    key = (
        os.environ.get("SUPABASE_KEY")
        or os.environ.get("SUPABASE_ANON_KEY")
        or os.environ.get("VITE_SUPABASE_ANON_KEY")
    )
    if not url or not key:
        raise ValueError("Supabase credentials are not set in environment variables.")
    client = create_client(url, key)
    if auth_header and auth_header.startswith("Bearer "):
        token = auth_header.split(" ", 1)[1]
        client.postgrest.auth(token)
    return client


# ─── Pydantic schemas ─────────────────────────────────────────────────────────
class QARequest(BaseModel):
    document_text: str
    question: str
    history: List[dict] = []
    tender_id: Optional[str] = None


class GoNoGoRequest(BaseModel):
    analysis_result: dict
    company_profile: str


class DraftProposalRequest(BaseModel):
    tender_id: str
    org_id: str
    custom_instructions: Optional[str] = ""
    reference_ids: Optional[List[str]] = []



# ─── Routes ───────────────────────────────────────────────────────────────────
@app.get("/")
def read_root():
    return {
        "name": "TenderIQ API",
        "version": "1.0.0",
        "description": "FastAPI backend for AI-powered Tender Analysis & Intelligence",
        "endpoints": {
            "health": "/api/health",
            "upload": "/api/upload",
            "qa": "/api/qa",
            "gonogo": "/api/gonogo",
        },
    }


@app.get("/api/health")
def health_check():
    return {
        "status": "healthy",
        "gemini_api_configured": os.environ.get("GEMINI_API_KEY") is not None,
        "supabase_configured": os.environ.get("SUPABASE_URL") is not None,
    }


@app.get("/api/tenders/{tender_id}/status")
def get_tender_status(
    tender_id: str, authorization: Optional[str] = Header(None)
):
    """Poll the processing status of a tender."""
    try:
        supabase = get_supabase(authorization)
        res = (
            supabase.table("tenders")
            .select("id,status,analysis_result,extracted_text,page_count,file_size,deadline,name,created_at")
            .eq("id", tender_id)
            .single()
            .execute()
        )
        return res.data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    tender_id: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    file: Optional[UploadFile] = File(None),
    raw_text: Optional[str] = Form(None),
    filename: Optional[str] = Form(None),
    authorization: Optional[str] = Header(None),
):
    """
    Async upload endpoint for Render persistent server.
    Reads file bytes immediately, then returns 202 and schedules
    all heavy processing (AI analysis + embedding) as a background task.
    The frontend polls /api/tenders/{id}/status to track progress.
    """
    if not gemini_client:
        raise HTTPException(
            status_code=500,
            detail="Gemini API Client is not configured. Check backend environment variables.",
        )

    # ── Read file bytes NOW (UploadFile objects become invalid after response) ─
    files_data: List[tuple] = []
    all_files = list(files or [])
    if file:
        all_files.append(file)

    for f in all_files:
        data = await f.read()
        files_data.append((data, f.filename))

    # ── Schedule background processing and return 202 immediately ─────────────
    background_tasks.add_task(
        process_tender_background,
        tender_id=tender_id,
        files_data=files_data,
        raw_text=raw_text,
        authorization=authorization,
    )

    return {
        "status": "queued",
        "tender_id": tender_id,
        "message": "Tender queued for background processing. Poll /api/tenders/{id}/status for updates.",
    }


def process_tender_background(
    tender_id: str,
    files_data: List[tuple],
    raw_text: Optional[str],
    authorization: Optional[str],
    is_lead: bool = False,
    lead_title: Optional[str] = None
):
    """
    Background worker — runs AFTER the 202 response is sent to the client.
    On Render (persistent server) this function runs to completion.
    """
    # ── 1. Text extraction ────────────────────────────────────────────────────
    extracted_text = ""
    page_count = 0
    file_size = 0

    try:
        if is_lead and gemini_client:
            print(f"[BG] Simulating full specification text expansion for imported lead {tender_id}...")
            # We use Gemini to generate a highly detailed tender spec so that the compliance audit is fully detailed.
            prompt = f"""
            You are a senior government procurement auditor. 
            Generate a highly detailed, professional, and comprehensive official Tender Specification Document based on this matching lead opportunity:
            
            Title: {lead_title or "Tender Opportunity"}
            Summary: {raw_text or ""}
            
            Your generated document MUST contain these exact section headings with rich, detailed specifications, clauses, numbers, and parameters:
            - SECTION 1: INTRODUCTION & EXECUTIVE SUMMARY
            - SECTION 2: TECHNICAL SPECIFICATIONS & WORK SCOPE (describe specific engineering standards, metrics, quantities)
            - SECTION 3: ELIGIBILITY & VENDOR QUALIFICATIONS (must mention ISO 9001/14001, turnover, minimum 5 years experience)
            - SECTION 4: KEY DATES & BID TIMELINE (include bid submission deadline, opening date, pre-bid meeting date in active future 2026/2027)
            - SECTION 5: FINANCIAL CLAUSES (list Earnest Money Deposit/EMD, turnover limits, payment schedules)
            - SECTION 6: REQUIRED DOCUMENTS FOR SUBMISSION (list specific forms, certifications, declarations, and balance sheets to include)
            - SECTION 7: RISKS, PENALTIES & LIQUIDATED DAMAGES (mention delay penalties, performance guarantees, and specifications compliance risks)
            - SECTION 8: EVALUATION METHODOLOGY (state selection criteria, technical vs financial weightage, L1 rules)
            - SECTION 9: TENDER CONTACT DETAILS & SUBMISSION FORUM
            
            Use official, formal procurement language. Make the document long, descriptive, and highly detailed (around 1500 to 2000 words) so that a parser can extract comprehensive details for every category.
            """
            response = gemini_client.model.generate_content(prompt)
            extracted_text = response.text.strip()
            page_count = max(4, len(extracted_text) // 2500)
            file_size = len(extracted_text.encode("utf-8"))
        elif files_data:
            parts, total_pages, total_size = [], 0, 0
            for data, fname in files_data:
                total_size += len(data)
                f_text, f_pages = extract_content(data, fname)
                parts.append(
                    f"\n=========================================\n"
                    f"DOCUMENT: {fname}\n"
                    f"=========================================\n\n{f_text}"
                )
                total_pages += f_pages
            extracted_text = "\n".join(parts)
            page_count = total_pages
            file_size = total_size

        elif raw_text:
            extracted_text = raw_text.strip()
            file_size = len(extracted_text.encode("utf-8"))
            page_count = max(1, len(extracted_text) // 3000 + 1)

        else:
            _mark_failed(tender_id, authorization)
            print(f"[BG] No files or raw_text provided for tender {tender_id}")
            return

    except Exception as e:
        _mark_failed(tender_id, authorization)
        print(f"[BG] Text extraction failed for tender {tender_id}: {e}")
        return

    # ── 2. AI compliance analysis ─────────────────────────────────────────────
    try:
        print(f"[BG] Starting AI analysis for tender {tender_id}...")
        analysis = gemini_client.extract_tender_details(extracted_text)
        deadline = analysis.get("deadline")
        print(f"[BG] AI analysis complete for tender {tender_id}")
    except Exception as e:
        _mark_failed(tender_id, authorization)
        print(f"[BG] AI analysis failed for tender {tender_id}: {e}")
        return

    # ── 3. Update tender record to Active ─────────────────────────────────────
    try:
        supabase = get_supabase(authorization)
        supabase.table("tenders").update(
            {
                "status": "Active",
                "page_count": page_count,
                "file_size": file_size,
                "extracted_text": extracted_text[:100_000],
                "analysis_result": analysis,
                "deadline": deadline,
            }
        ).eq("id", tender_id).execute()
        print(f"[BG] Tender {tender_id} updated to Active")
    except Exception as e:
        print(f"[BG] DB update failed for tender {tender_id}: {e}")
        return

    # ── 4. RAG chunking & embeddings ─────────────────────────────────────────
    try:
        print(f"[BG] Starting RAG chunking for tender {tender_id}...")
        chunks = gemini_client.chunk_text(extracted_text)
        if chunks:
            print(f"[BG] Generating embeddings for {len(chunks)} chunks...")
            embeddings = gemini_client.generate_embeddings(chunks)

            # Delete old chunks for this tender before inserting fresh ones
            supabase.table("tender_chunks").delete().eq("tender_id", tender_id).execute()

            payload = [
                {
                    "tender_id": tender_id,
                    "chunk_content": chunk,
                    "page_number": idx + 1,
                    "embedding": emb,
                }
                for idx, (chunk, emb) in enumerate(zip(chunks, embeddings))
            ]
            if payload:
                supabase.table("tender_chunks").insert(payload).execute()
                print(f"[BG] Inserted {len(payload)} chunks for tender {tender_id}")
    except Exception as e:
        # Chunking failure is non-fatal — tender is already Active
        print(f"[BG] RAG chunking skipped for tender {tender_id}: {e}")


def _mark_failed(tender_id: str, auth_header: Optional[str]):
    """Helper to update a tender row to Failed status on error."""
    try:
        supabase = get_supabase(auth_header)
        supabase.table("tenders").update({"status": "Failed"}).eq("id", tender_id).execute()
    except Exception as db_err:
        print(f"Could not mark tender {tender_id} as Failed: {db_err}")


@app.post("/api/qa")
def ask_question(
    request: QARequest, authorization: Optional[str] = Header(None)
):
    """
    Answers a free-form user query.
    Primary context: analysis_result (structured AI extraction) fetched from DB.
    Secondary context: pgvector RAG chunks.
    Final fallback: raw document_text from request body.
    """
    if not gemini_client:
        raise HTTPException(status_code=500, detail="Gemini API Client is not configured.")

    rag_context = ""
    analysis_result = None

    if request.tender_id:
        try:
            supabase = get_supabase(authorization)

            # Fetch structured analysis_result
            tender_row = (
                supabase.table("tenders")
                .select("analysis_result")
                .eq("id", request.tender_id)
                .single()
                .execute()
            )
            if tender_row.data:
                analysis_result = tender_row.data.get("analysis_result")

            # RAG vector search for supplementary chunks
            try:
                query_embeddings = gemini_client.generate_embeddings([request.question])
                if query_embeddings:
                    rpc_res = supabase.rpc(
                        "match_tender_chunks",
                        {
                            "query_embedding": query_embeddings[0],
                            "match_threshold": 0.25,
                            "match_count": 5,
                            "filter_tender_id": request.tender_id,
                        },
                    ).execute()
                    if rpc_res.data:
                        rag_context = "\n\n".join(
                            row["chunk_content"] for row in rpc_res.data
                        )
            except Exception as rag_err:
                print(f"RAG search skipped: {rag_err}")

        except Exception as err:
            print(f"Could not fetch tender data for QA: {err}")

    raw_context = rag_context or request.document_text[:80_000]

    try:
        answer = gemini_client.ask_tender_question(
            document_text=raw_context,
            question=request.question,
            history=request.history,
            analysis_result=analysis_result,
        )
        return {"answer": answer}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/gonogo")
def evaluate_bidding_suitability(request: GoNoGoRequest):
    """Analyses a company's suitability (0-100 score) against the tender."""
    if not gemini_client:
        raise HTTPException(status_code=500, detail="Gemini API Client is not configured.")
    try:
        evaluation = gemini_client.evaluate_go_nogo(
            analysis_result=request.analysis_result,
            company_profile=request.company_profile,
        )
        return evaluation
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Reference Library & Proposal Writer Endpoints ────────────────────────────

@app.post("/api/references/upload")
async def upload_reference_document(
    org_id: str = Form(...),
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(None)
):
    """Uploads a reference document, extracts text content, and saves it to workspace_references."""
    try:
        file_bytes = await file.read()
        text_content, _ = extract_content(file_bytes, file.filename)
        
        db = get_supabase(authorization)
        
        # Get active user ID
        user_id = None
        try:
            user_res = db.auth.get_user()
            if user_res and user_res.user:
                user_id = user_res.user.id
        except Exception:
            pass

        res = db.table("workspace_references").insert({
            "org_id": org_id,
            "user_id": user_id,
            "filename": file.filename,
            "file_size": len(file_bytes),
            "content_text": text_content
        }).execute()

        if not res.data:
            raise HTTPException(status_code=400, detail="Failed to save reference document.")

        return res.data[0]
    except ExtractionError as ee:
        raise HTTPException(status_code=400, detail=str(ee))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def create_docx_file(tender_name: str, draft: dict) -> io.BytesIO:
    """Helper to generate a beautifully styled MS Word Document response."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def add_runs(paragraph, text, font_name="Arial", font_size_pt=11, color_rgb=None, bold=False, italic=False):
        """Helper to parse **bold** text and add formatted runs to a paragraph."""
        parts = text.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            if color_rgb:
                run.font.color.rgb = color_rgb
            run.font.bold = bold or (i % 2 == 1)
            run.font.italic = italic

    doc = Document()
    
    # 1. Title Page / Cover Page
    for _ in range(3):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(title_p, "BID PROPOSAL RESPONSE", font_size_pt=28, color_rgb=RGBColor(30, 58, 138), bold=True)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(sub_p, f"Project/Tender: {tender_name}", font_size_pt=14, color_rgb=RGBColor(100, 116, 139), italic=True)
    
    for _ in range(4):
        doc.add_paragraph()
        
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(meta_p, "Prepared by: Bidding Team\nGenerated via TenderIQ Proposal Engine", color_rgb=RGBColor(71, 85, 105))
    
    doc.add_page_break()
    
    # 2. Executive Cover Letter
    h1 = doc.add_paragraph()
    add_runs(h1, "SECTION 1: Executive Cover Letter", font_size_pt=18, color_rgb=RGBColor(30, 58, 138), bold=True)
    
    cover_text = draft.get("cover_letter", "")
    for para in cover_text.split("\n\n"):
        if para.strip():
            p = doc.add_paragraph()
            add_runs(p, para.strip())
            p.paragraph_format.line_spacing = 1.15
            
    doc.add_page_break()
    
    # 3. Technical Response Section
    h2 = doc.add_paragraph()
    add_runs(h2, "SECTION 2: Technical Response & Statement of Work", font_size_pt=18, color_rgb=RGBColor(30, 58, 138), bold=True)
    
    tech_text = draft.get("technical_response", "")
    for para in tech_text.split("\n\n"):
        para_strip = para.strip()
        if not para_strip:
            continue
        if para_strip.startswith("### "):
            p = doc.add_paragraph()
            add_runs(p, para_strip[4:], font_size_pt=13, color_rgb=RGBColor(15, 118, 110), bold=True)
        elif para_strip.startswith("## "):
            p = doc.add_paragraph()
            add_runs(p, para_strip[3:], font_size_pt=15, color_rgb=RGBColor(30, 58, 138), bold=True)
        else:
            lines = para_strip.split("\n")
            if len(lines) > 1 and all(l.strip().startswith("-") or l.strip().startswith("*") for l in lines if l.strip()):
                for line in lines:
                    if line.strip():
                        clean_line = line.strip()[1:].strip()
                        bp = doc.add_paragraph(style='List Bullet')
                        add_runs(bp, clean_line)
            else:
                p = doc.add_paragraph()
                add_runs(p, para_strip)
                p.paragraph_format.line_spacing = 1.15
                
    doc.add_page_break()
    
    # 4. Capability Compliance Matrix
    h3 = doc.add_paragraph()
    add_runs(h3, "SECTION 3: Capability Compliance Matrix", font_size_pt=18, color_rgb=RGBColor(30, 58, 138), bold=True)
    
    matrix = draft.get("capability_matrix", [])
    if matrix:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        add_runs(hdr_cells[0].paragraphs[0], 'Tender Requirement', font_size_pt=10.5, bold=True)
        add_runs(hdr_cells[1].paragraphs[0], 'Compliance Status', font_size_pt=10.5, bold=True)
        add_runs(hdr_cells[2].paragraphs[0], 'Evidence / Reference', font_size_pt=10.5, bold=True)
        
        for item in matrix:
            row_cells = table.add_row().cells
            add_runs(row_cells[0].paragraphs[0], item.get("requirement", "N/A"), font_size_pt=10)
            add_runs(row_cells[1].paragraphs[0], item.get("compliance_status", "Compliant"), font_size_pt=10)
            add_runs(row_cells[2].paragraphs[0], item.get("evidence_reference", "N/A"), font_size_pt=10)
    else:
        p = doc.add_paragraph()
        add_runs(p, "No compliance requirements mapped.", italic=True)
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


class DownloadProposalRequest(BaseModel):
    tender_name: str
    draft: dict


@app.post("/api/proposal/draft")
def draft_proposal_response(
    request: DraftProposalRequest,
    authorization: Optional[str] = Header(None)
):
    """Generates a structured proposal draft from references and returns the JSON draft data."""
    if not gemini_client:
        raise HTTPException(status_code=500, detail="Gemini API Client is not configured.")
    try:
        db = get_supabase(authorization)
        
        # 1. Fetch active tender metadata and analysis
        tender_res = db.table("tenders").select("name, analysis_result").eq("id", request.tender_id).single().execute()
        if not tender_res.data:
            raise HTTPException(status_code=404, detail="Tender not found or access denied.")
        tender = tender_res.data

        # 2. Fetch selected reference library records
        query = db.table("workspace_references").select("filename, content_text").eq("org_id", request.org_id)
        if request.reference_ids:
            query = query.in_("id", request.reference_ids)
        refs_res = query.execute()
        references = refs_res.data or []

        # 3. Call Gemini to draft the content
        draft_data = gemini_client.draft_proposal(
            tender_name=tender["name"],
            tender_analysis=tender["analysis_result"],
            references=references,
            custom_instructions=request.custom_instructions
        )

        return draft_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/proposal/download")
def download_proposal_document(request: DownloadProposalRequest):
    """Generates a styled DOCX file from the preview draft and returns the file stream."""
    try:
        file_stream = create_docx_file(request.tender_name, request.draft)
        filename = f"{request.tender_name.replace(' ', '_')}_Bid_Proposal.docx"
        
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class AudioBriefRequest(BaseModel):
    tender_id: str


@app.post("/api/tenders/audio-brief")
def get_audio_briefing(
    request: AudioBriefRequest,
    authorization: Optional[str] = Header(None)
):
    """Generates a conversational podcast script and returns it as a streaming MP3 response."""
    from fastapi.responses import StreamingResponse
    from gtts import gTTS
    import io

    if not gemini_client:
        raise HTTPException(status_code=500, detail="Gemini API Client is not configured.")
    
    try:
        db = get_supabase(authorization)
        
        # Fetch tender name and analysis
        tender_res = db.table("tenders").select("name, analysis_result").eq("id", request.tender_id).single().execute()
        if not tender_res.data:
            raise HTTPException(status_code=404, detail="Tender not found or access denied.")
        tender = tender_res.data
        
        name = tender.get("name", "Active Tender")
        analysis = tender.get("analysis_result", {})
        
        # Generate script using Gemini
        script = gemini_client.generate_audio_briefing_script(name, analysis)
        
        # Convert script to MP3 bytes using gTTS
        tts = gTTS(text=script, lang='en', tld='com')
        audio_stream = io.BytesIO()
        tts.write_to_fp(audio_stream)
        audio_stream.seek(0)
        
        filename = f"briefing_{request.tender_id}.mp3"
        return StreamingResponse(
            audio_stream,
            media_type="audio/mpeg",
            headers={"Content-Disposition": f"inline; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── LEAD GENERATION / AUTOMATED RFP TRACKING SCHEMAS & ENDPOINTS ────────────

class SettingsRequest(BaseModel):
    org_id: str
    company_profile: str
    alert_keywords: str
    slack_webhook: Optional[str] = ""


class CrawlRequest(BaseModel):
    org_id: str


@app.get("/api/leads/settings")
def get_leads_settings(
    org_id: str,
    authorization: Optional[str] = Header(None)
):
    """Fetches organization matching settings."""
    try:
        db = get_supabase(authorization)
        res = db.table("organizations").select("company_profile, alert_keywords, slack_webhook").eq("id", org_id).single().execute()
        return res.data or {"company_profile": "", "alert_keywords": "", "slack_webhook": ""}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/leads/settings")
def save_leads_settings(
    request: SettingsRequest,
    authorization: Optional[str] = Header(None)
):
    """Saves organization matching settings."""
    try:
        db = get_supabase(authorization)
        res = db.table("organizations").update({
            "company_profile": request.company_profile,
            "alert_keywords": request.alert_keywords,
            "slack_webhook": request.slack_webhook
        }).eq("id", request.org_id).execute()
        return {"status": "success", "data": res.data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/leads")
def get_crawled_leads(
    org_id: str,
    authorization: Optional[str] = Header(None)
):
    """Lists crawled RFP leads for the organization."""
    try:
        db = get_supabase(authorization)
        res = db.table("crawled_tenders").select("*").eq("org_id", org_id).order("created_at", desc=True).execute()
        return res.data or []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/leads/crawl")
def trigger_leads_crawl(
    request: CrawlRequest,
    authorization: Optional[str] = Header(None)
):
    """Simulates crawling government portals, evaluates suitability, and alerts Slack."""
    from crawler import simulate_government_portal_crawl
    from matcher import match_tenders_to_profile
    from notifier import send_slack_alert

    try:
        db = get_supabase(authorization)
        
        # 1. Fetch organization settings and details
        org_res = db.table("organizations").select("name, company_profile, alert_keywords, slack_webhook").eq("id", request.org_id).single().execute()
        if not org_res.data:
            raise HTTPException(status_code=404, detail="Organization not found.")
        org = org_res.data
        
        keywords_str = org.get("alert_keywords") or ""
        profile = org.get("company_profile") or ""
        slack_url = org.get("slack_webhook") or ""
        org_name = org.get("name") or "Your Organization"
        
        keywords = [k.strip() for k in keywords_str.split(",") if k.strip()]
        
        # 2. Run crawler simulation
        crawled_items = simulate_government_portal_crawl(keywords)
        
        # 3. Match suitability
        matches = match_tenders_to_profile(crawled_items, profile)
        
        # 4. Insert matched results into DB
        inserted_matches = []
        for m in matches:
            # Check if this lead already exists for this org
            existing = db.table("crawled_tenders").select("id").eq("org_id", request.org_id).eq("title", m["title"]).execute()
            if existing.data:
                continue
                
            insert_data = {
                "org_id": request.org_id,
                "title": m["title"],
                "portal_name": m["portal_name"],
                "tender_value": m["tender_value"],
                "deadline": m["deadline"],
                "description": m["description"],
                "compatibility_score": m["compatibility_score"],
                "compatibility_reason": m["compatibility_reason"],
                "imported": False
            }
            res = db.table("crawled_tenders").insert(insert_data).execute()
            if res.data:
                inserted_matches.append(res.data[0])
                
        # 5. Dispatch Slack notifications for high matches
        if inserted_matches and slack_url:
            send_slack_alert(slack_url, org_name, inserted_matches)
            
        return {
            "status": "success",
            "new_matches_count": len(inserted_matches),
            "matches": inserted_matches
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def get_user_id_from_token(auth_header: Optional[str]) -> Optional[str]:
    import base64
    import json
    if not auth_header or not auth_header.startswith("Bearer "):
        return None
    try:
        token = auth_header.split(" ", 1)[1]
        parts = token.split(".")
        if len(parts) != 3:
            return None
        payload_b64 = parts[1]
        padding = "=" * (4 - len(payload_b64) % 4)
        payload_bytes = base64.b64decode(payload_b64 + padding)
        payload_data = json.loads(payload_bytes)
        return payload_data.get("sub")
    except Exception as e:
        print(f"Error parsing token: {e}")
        return None


@app.post("/api/leads/{lead_id}/import")
def import_lead_to_workspace(
    lead_id: str,
    background_tasks: BackgroundTasks,
    authorization: Optional[str] = Header(None)
):
    """Imports a crawled lead as an active analyzed tender in the workspace."""
    try:
        db = get_supabase(authorization)
        lead_res = db.table("crawled_tenders").select("*").eq("id", lead_id).single().execute()
        if not lead_res.data:
            raise HTTPException(status_code=404, detail="Lead not found.")
        lead = lead_res.data
        
        if lead.get("imported"):
            raise HTTPException(status_code=400, detail="Lead already imported.")
            
        # Get active user id from JWT token to satisfy RLS check
        user_id = get_user_id_from_token(authorization)
        
        # Insert new tender record
        tender_res = db.table("tenders").insert({
            "user_id": user_id,
            "org_id": lead.get("org_id"),
            "name": lead.get("title"),
            "status": "Processing",
            "file_size": len(lead.get("description", "").encode("utf-8")),
            "page_count": 1,
            "kanban_stage": "Discovered"
        }).execute()
        
        if not tender_res.data:
            raise HTTPException(status_code=500, detail="Failed to create tender.")
            
        new_tender = tender_res.data[0]
        new_tender_id = new_tender.get("id")
        
        # Schedule background tender analysis
        background_tasks.add_task(
            process_tender_background,
            tender_id=new_tender_id,
            files_data=None,
            raw_text=lead.get("description", ""),
            authorization=authorization,
            is_lead=True,
            lead_title=lead.get("title")
        )
        
        # Mark lead as imported
        db.table("crawled_tenders").update({"imported": True}).eq("id", lead_id).execute()
        
        return {
            "status": "success",
            "tender_id": new_tender_id,
            "message": "Lead imported successfully. Running background compliance analysis."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    import io
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)


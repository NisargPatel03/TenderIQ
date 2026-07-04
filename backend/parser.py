import io
import fitz  # PyMuPDF
from docx import Document

class ExtractionError(Exception):
    """Custom exception raised for file parsing or extraction errors."""
    pass

def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extracts text and page count from PDF bytes using PyMuPDF (fitz)."""
    try:
        # Open PDF from bytes memory buffer
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(pdf_document)
        
        full_text = []
        for page_num in range(page_count):
            page = pdf_document.load_page(page_num)
            text = page.get_text()
            full_text.append(text)
            
        pdf_document.close()
        
        extracted_text = "\n".join(full_text).strip()
        
        # Check if the extracted text is empty or extremely short (scanned PDF detection)
        if len(extracted_text) < 50:
            raise ExtractionError(
                "This PDF document appears to be scanned or contains image-only pages. "
                "TenderIQ requires a digital, text-based PDF to extract information."
            )
            
        return extracted_text, page_count
    except Exception as e:
        if isinstance(e, ExtractionError):
            raise e
        raise ExtractionError(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extracts text and page count estimation from DOCX bytes using python-docx."""
    try:
        doc_io = io.BytesIO(file_bytes)
        doc = Document(doc_io)
        
        # Extract text from paragraphs
        paragraphs_text = [p.text for p in doc.paragraphs]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_cells_text))
                
        full_text = "\n".join(paragraphs_text + table_text).strip()
        
        # Estimate page count based on ~500 words per page
        word_count = len(full_text.split())
        estimated_pages = max(1, (word_count // 500) + 1)
        
        if len(full_text) < 10:
            raise ExtractionError("The Word document is empty or could not be read.")
            
        return full_text, estimated_pages
    except Exception as e:
        if isinstance(e, ExtractionError):
            raise e
        raise ExtractionError(f"Failed to extract text from Word document: {str(e)}")

def extract_text_from_txt(file_bytes: bytes) -> tuple[str, int]:
    """Extracts text and page count estimation from TXT bytes."""
    try:
        try:
            text = file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1").strip()
            
        if not text:
            raise ExtractionError("The text file is empty.")
            
        # Estimate pages based on characters (~3000 chars per page)
        char_count = len(text)
        estimated_pages = max(1, (char_count // 3000) + 1)
        
        return text, estimated_pages
    except Exception as e:
        if isinstance(e, ExtractionError):
            raise e
        raise ExtractionError(f"Failed to read text file: {str(e)}")

def extract_content(file_bytes: bytes, filename: str) -> tuple[str, int]:
    """Main function to select parser based on file extension and extract text and page count."""
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext in ["txt"]:
        return extract_text_from_txt(file_bytes)
    else:
        raise ExtractionError(f"Unsupported file format: .{ext}")
